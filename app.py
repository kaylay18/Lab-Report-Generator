import gradio as gr
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email import encoders

def create_report(name, prof, course, date, target_email, file):
  data = pd.read_csv(file.name)
  doc = Document()
  doc.add_paragraph(name)
  doc.add_paragraph(prof)
  doc.add_paragraph(course)
  doc.add_paragraph(date)
  doc.add_heading('Fluid Dynamics Experiment Report', level = 1)

  doc.add_heading('Introduction', level = 2)
  doc.add_paragraph("In fluid dynamics, understanding the behavior of fluids in motion is essential for designing and optimizing systems such as pipelines, pumps, and valves in various engineering applications. Collecting data on key parameters enables engineers to predict fluid behavior, optimize system performance, and troubleshoot inefficiencies.")
  doc.add_paragraph("Flow rate (Q), the volume of fluid passing through a pipe per unit of time, is crucial for determining system capacity and efficiency, as it directly influences parameters like pressure drop and velocity. Pipe diameter (D), the internal width of the pipe, significantly affects both flow rate and fluid velocity, with larger diameters reducing resistance and enabling higher flow rates. Fluid velocity (v), the speed at which the fluid moves, is important for maintaining desired flow regimes and preventing issues like erosion or blockages. Pressure drop (ΔP), the difference in pressure between two points in a pipe, reflects the resistance to fluid flow and helps identify energy losses, ensuring efficient system performance. Finally, fluid density (ρ), the mass of fluid per unit volume, influences how fluids behave under various conditions, with denser fluids generating higher pressure drops and requiring more energy to pump. Collecting data on these variables allows engineers to design more efficient, reliable, and safe fluid transport systems while minimizing energy consumption and operational costs.")
  doc.add_heading('Summary Statistics', level = 2)
  summary_stats = data.describe()
  table = doc.add_table(rows=1, cols=len(summary_stats.columns) + 1)
  hdr_cells = table.rows[0].cells

  # Add column headers
  hdr_cells[0].text = 'Statistic'  # Column for statistic names
  for i, col in enumerate(summary_stats.columns):
      hdr_cells[i + 1].text = col

  # Add rows for each statistic
  for stat_name, row in summary_stats.iterrows():
      row_cells = table.add_row().cells
      row_cells[0].text = stat_name  # Stat name in the first column
      for i, val in enumerate(row):
          row_cells[i + 1].text = str(val)  # Values in subsequent columns

   # Plotting Flow Rate vs. Fluid Velocity
  plt.figure(figsize=(10, 6))
  plt.plot(data['Fluid Velocity (v) [m/s]'], data['Flow Rate (Q) [m^3/s]'], marker='o', linestyle='None', color='blue', label='Flow Rate vs. Fluid Velocity')
  plt.title('Flow Rate vs. Fluid Velocity')
  plt.xlabel('Fluid Velocity (m/s)')
  plt.ylabel('Flow Rate (m³/s)')
  plt.grid()
  plt.axhline(0, color='black', linewidth=0.5, ls='--')
  plt.axvline(0, color='black', linewidth=0.5, ls='--')
  plt.legend()
  plt.savefig('flow_rate_vs_velocity.png')  # Save the plot
  plt.close()
    # Add the plot image to the Word document
  # doc.add_page_break()
  doc.add_heading('Flow Rate vs. Fluid Velocity:', level=2)
  doc.add_paragraph('Plot:')
  doc.add_picture('flow_rate_vs_velocity.png', width=Inches(6.0))
  doc.add_paragraph("For a given flow pipe diameter, flow rate (Q) is related to velocity (v) by the equation:")
  doc.add_paragraph("Q = v * π * D^2 / 4").alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("where A is the cross-sectional area of the pipe and D is the pipe diameter. The expected behavior is that flow rate increases linearly with fluid velocity. As the diameter changes, this linear relationship still holds and the slope will vary depending on the pipe's size.")
  doc.add_paragraph("This relationship implies that, for a given pipe diameter, the flow rate will increase linearly with velocity. However, real-world factors such as turbulence, flow regime changes, and pipe roughness can affect this proportionality. For instance, at higher velocities, turbulence could disrupt the smooth flow, leading to non-linear increases in flow rate. Thus, deviations from this expected relationship might occur, especially in systems operating at high Reynolds numbers (indicative of turbulent flow).")
  doc.add_paragraph("During experimentation, it’s important to monitor the Reynolds number and flow regime (laminar or turbulent), as this can influence the linearity of the flow rate-velocity relationship. In turbulent flow, additional losses and irregularities in the velocity profile may cause the flow rate to increase more slowly than expected as velocity increases.")
  
  # Plotting Flow Rate vs. Pressure Drop
  plt.figure(figsize=(10, 6))
  plt.plot(data['Pressure Drop (ΔP) [Pa]'], data['Flow Rate (Q) [m^3/s]'], marker='o', linestyle='None', color='blue', label='Flow Rate vs. Pressure Drop')
  plt.title('Flow Rate vs. Pressure Drop')
  plt.xlabel('Pressure Drop (Pa)')
  plt.ylabel('Flow Rate (m³/s)')
  plt.grid()
  plt.axhline(0, color='black', linewidth=0.5, ls='--')
  plt.axvline(0, color='black', linewidth=0.5, ls='--')
  plt.legend()
  plt.savefig('flow_rate_vs_pressure_drop.png')  # Save the plot
  plt.close()
    # Add the plot image to the Word document
  # doc.add_page_break()
  doc.add_heading('Flow Rate vs. Pressure Drop:', level=2)
  doc.add_paragraph('Plot:')
  doc.add_picture('flow_rate_vs_pressure_drop.png', width=Inches(6.0))
  doc.add_paragraph("The pressure drop along a pipe is governed by the Darcy-Weisbach equation:")
  doc.add_paragraph("ΔP = f * L/D * ρv^2/2").alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("where f is the fanning friction factor which depends on flow regime and pipe roughness, L is the pipe length, ρ is fluid density,and v is fluid velocity. For turbulent flow (Re > 4000), pressure drop should increase quadratically with velocity and hence with flow rate. For laminar flow (Re < 2000), the relationship is linear.")
  doc.add_paragraph("Since velocity (v) and flow rate (Q) are related, the pressure drop is expected to increase quadratically with flow rate in turbulent conditions (due to the v^2 term). In laminar flow, the relationship between pressure drop and flow rate is linear. However, real systems may experience variations in pipe friction factor (f), influenced by factors such as roughness, fouling, or changes in flow regime, which can cause deviations from the predicted relationship.")
  doc.add_paragraph("Observing the pressure drop across different flow rates helps identify whether the system is operating in laminar or turbulent flow regimes. If the pressure drop does not follow the expected quadratic relationship in turbulent flow, it may indicate additional frictional losses or changes in pipe surface conditions that should be investigated.")

  # Plotting Pipe Diameter vs. Pressure Drop
  plt.figure(figsize=(10, 6))
  plt.plot(data['Pressure Drop (ΔP) [Pa]'], data['Pipe Diameter (D) [m]'], marker='o', linestyle='None', color='blue', label='Pipe Diameter vs. Pressure Drop')
  plt.title('Pipe Diameter vs. Pressure Drop')
  plt.xlabel('Pressure Drop (Pa)')
  plt.ylabel('Pipe Diameter (m)')
  plt.grid()
  plt.axhline(0, color='black', linewidth=0.5, ls='--')
  plt.axvline(0, color='black', linewidth=0.5, ls='--')
  plt.legend()
  plt.savefig('pressure_drop_vs_pipe_diam.png')  # Save the plot
  plt.close()
    # Add the plot image to the Word document
  # doc.add_page_break()
  doc.add_heading('Pressure Drop vs. Pipe Diameter:', level=2)
  doc.add_paragraph('Plot:')
  doc.add_picture('pressure_drop_vs_pipe_diam.png', width=Inches(6.0))
  doc.add_paragraph("From the Darcy-Weisbach equation, pressure drop ΔP is inversely proportional to the pipe diameter:")
  doc.add_paragraph("ΔP ∝ 1/D^5").alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("The expected behavior from this relationship is that as the pipe diameter increases, the pressure drop decreases significantly. This is due to the larger cross-sectional area reducing the resistance to flow.")
  doc.add_paragraph("Accurate measurements of pipe diameter are critical, as small errors can disproportionately affect the pressure drop data. Additionally, bends or elbows in the piping system can introduce pressure losses that are not accounted for by simple diameter changes. These factors should be considered in experiments to align the results more closely with theoretical expectations.")

  # Plotting Pipe Diameter vs. Flow Rate
  plt.figure(figsize=(10, 6))
  plt.plot(data['Flow Rate (Q) [m^3/s]'], data['Pipe Diameter (D) [m]'], marker='o', linestyle='None', color='blue', label='Flow Rate vs. Pipe Diameter')
  plt.title('Pipe Diameter vs. Flow Rate')
  plt.xlabel('Flow Rate (m³/s)')
  plt.ylabel('Pipe Diameter (m)')
  plt.grid()
  plt.axhline(0, color='black', linewidth=0.5, ls='--')
  plt.axvline(0, color='black', linewidth=0.5, ls='--')
  plt.legend()
  plt.savefig('flow_rate_vs_pipe_diam.png')  # Save the plot
  plt.close()
    # Add the plot image to the Word document
  # doc.add_page_break()
  doc.add_heading('Flow Rate vs. Pipe Diameter:', level=2)
  doc.add_paragraph('Plot:')
  doc.add_picture('flow_rate_vs_pipe_diam.png', width=Inches(6.0))
  doc.add_paragraph("For a given velocity, flow rate is proportional to the square of the diameter:")
  doc.add_paragraph("Q ∝ D^2").alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("A larger pipe diameter allows more fluid to pass through, increasing the flow rate. This is a strong dependency, so even a small increase in diameter can significantly raise flow rate. In real systems, this relationship may be affected by factors such as the flow regime, pipe fittings, and changes in fluid properties (such as density or viscosity) that could alter the velocity profile within the pipe.")
  doc.add_paragraph("In experimentation, deviations from the expected relationship could indicate changes in flow velocity that are not proportional to diameter or the presence of turbulent effects that disrupt the smooth relationship. Additionally, care should be taken to ensure that measurements of diameter are accurate, as even slight errors can lead to significant changes in flow rate predictions.")

  # Plotting Pipe Diameter vs. Fluid Density
  plt.figure(figsize=(10, 6))
  plt.plot(data['Pressure Drop (ΔP) [Pa]'], data['Fluid Density (ρ) [kg/m^3]'], marker='o', linestyle='None', color='blue', label='Fluid Density vs. Pressure Drop')
  plt.title('Fluid Density vs. Pressure Drop')
  plt.xlabel('Pressure Drop (Pa)')
  plt.ylabel('Fluid Density (kg/m³)')
  plt.grid()
  plt.axhline(0, color='black', linewidth=0.5, ls='--')
  plt.axvline(0, color='black', linewidth=0.5, ls='--')
  plt.legend()
  plt.savefig('pressure_drop_vs_fluid_dens.png')  # Save the plot
  plt.close()
    # Add the plot image to the Word document
  doc.add_page_break()
  doc.add_heading('Pressure Drop vs. Fluid Density:', level=2)
  doc.add_paragraph('Plot:')
  doc.add_picture('pressure_drop_vs_fluid_dens.png', width=Inches(6.0))
  doc.add_paragraph("In the Darcy-Weisbach equation, pressure drop is directly proportional to fluid density:")
  doc.add_paragraph("ΔP ∝ ρ").alignment = WD_ALIGN_PARAGRAPH.CENTER
  doc.add_paragraph("This implies that denser fluids experience greater resistance to flow, resulting in a higher pressure drop for the same velocity. In real-world systems, fluctuations in temperature or the composition of the fluid can alter its density, leading to variations in pressure drop that may not be predicted by simple models.")
  doc.add_paragraph("In experiments, changes in temperature or composition (e.g., the introduction of impurities) can affect fluid density, causing the actual pressure drop to deviate from theoretical predictions. Monitoring and controlling fluid properties, such as temperature, can help ensure that the relationship between density and pressure drop remains consistent and predictable.")

  # doc.add_page_break()
  doc.add_heading('Analysis', level = 2)
  doc.add_paragraph("Discrepancies between the expected and actual results in fluid dynamics experiments can arise due to several factors. One common source of deviation is the assumption of ideal conditions, such as smooth pipe surfaces and fully developed flow, which may not hold true in real-world systems. Pipe roughness or fouling can increase resistance, leading to a higher-than-expected pressure drop for a given flow rate. Similarly, turbulence and other flow irregularities, especially at higher velocities, can cause deviations from theoretical relationships, as ideal models often assume laminar or steady flow. Inaccuracies in measuring variables like pipe diameter or fluid velocity can also introduce errors, particularly in calculating flow rates and pressure drops. Variations in fluid properties, such as non-constant density due to temperature changes, can further skew the results. Additionally, real systems may experience leaks or energy losses that are not accounted for in theoretical models, resulting in discrepancies between expected and measured outcomes. These factors highlight the importance of accounting for practical considerations and measurement errors when interpreting experimental data in fluid dynamics.")

  # doc.add_page_break()
  doc.add_heading('Conclusion', level = 2)
  doc.add_paragraph("Based on the expected relationships between the collected variables in this fluid dynamics experiment, several conclusions can be drawn. The data is likely to show that flow rate (Q) increases with fluid velocity (v), especially for a constant pipe diameter, due to the direct relationship between these two variables. As pipe diameter increases, both flow rate and velocity are expected to increase, reflecting the greater cross-sectional area available for fluid flow. The pressure drop (ΔP) is expected to decrease with increasing pipe diameter, as larger diameters reduce flow resistance, while it should increase with higher flow rates and velocities due to greater frictional losses within the pipe. Additionally, denser fluids will likely produce higher pressure drops under the same flow conditions, indicating the influence of fluid density (ρ) on system resistance. These relationships will validate fluid dynamic principles like the Darcy-Weisbach equation and allow for system optimization by carefully balancing flow rate, pipe diameter, and pressure drop. Understanding these interdependencies is critical for designing fluid transport systems that operate efficiently, safely, and cost-effectively.")

  # doc.add_page_break()
  doc.add_heading('Recommendations', level = 2)
  doc.add_paragraph("To improve data collection and reduce discrepancies in the experiment, several recommendations can be implemented. First, ensuring precise measurement instruments for flow rate, velocity, pressure drop, and pipe diameter is crucial for minimizing errors. Using calibrated sensors and regularly maintaining equipment will help achieve more accurate results. Minimizing pipe roughness and ensuring that the fluid flows through clean, well-maintained pipes can reduce resistance variations, leading to more consistent pressure drop measurements. Another recommendation is to control environmental factors, such as temperature, which can affect fluid density and viscosity, introducing variability into the results. Employing more frequent sampling and using automated data logging systems can capture more detailed variations over time, helping to identify outliers and patterns more effectively. Finally, conducting multiple trials under the same conditions and averaging the results can reduce the impact of random measurement errors and improve the reliability of the data. For future steps, further experimentation could explore how flow regime transitions, such as from laminar to turbulent flow, affect the relationships between variables. Introducing controlled variations in fluid properties, such as changing the fluid's temperature or composition, could help refine the understanding of how fluid density influences pressure drop and flow rate. Additionally, future experiments could investigate the impact of pipe length and roughness on system performance to better model real-world applications. Incorporating computational fluid dynamics (CFD) simulations alongside physical experiments could also provide deeper insights by allowing for virtual exploration of complex flow conditions that are difficult to replicate in the lab. By combining these approaches, future work can improve the accuracy of predictions and further optimize fluid transport system design.")

  output_file = 'data_report.docx'
  doc.save(output_file)

  send_email_with_attachment(target_email, name + ' - Final Lab Report',
                                "Hello Professor " + prof + ", hope you are doing well. Attached is my final lab report. Thank you and have a great rest of your day! -"+ name, output_file)

  return output_file

def send_email_with_attachment(recipient_email, subject, body, attachment_path):
    sender_email = "kfy001@bucknell.edu"  # Replace with your email
    sender_password = "vtjc fvka baio gzvt"  # Replace with your email password

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject
    msg.attach(MIMEText(body, 'plain'))

    # Attach the document
    attachment = open(attachment_path, 'rb')
    part = MIMEBase('application', 'octet-stream')
    part.set_payload(attachment.read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', f'attachment; filename={attachment_path}')
    msg.attach(part)

    try:
        server = smtplib.SMTP('smtp.gmail.com', 587)  # Use your SMTP server and port
        server.starttls()
        server.login(sender_email, sender_password)
        server.send_message(msg)
        print("Email sent successfully.")
    except Exception as e:
        print(f"Failed to send email: {e}")
    finally:
        server.quit()

def create_gradio_interface():
    interface = gr.Interface(
        fn=create_report,
        inputs=[
            gr.Textbox(label="Name", placeholder="Your Name Here"),
            gr.Textbox(label="Professor", placeholder="Professor Name Here"),
            gr.Textbox(label="Course", placeholder="Course Title Here"),
            gr.Textbox(label="Date", placeholder="Due Date Here"),
            gr.Textbox(label="Professor Email", placeholder="professor@email.com"),
            gr.File(label="Upload your CSV file")
        ],
        theme="glass",
        outputs="file",  # Output is a file (Word document)
        title="Data Analysis Report Generator for Fluid Dynamics Experiment",
        description="Upload a CSV file to generate a data analysis report as a Word document."
    )

    return interface

# Launch the Gradio interface
if __name__ == "__main__":
    create_gradio_interface().launch(debug=True)
